from __future__ import annotations

import argparse
import json
import math
import re
import shutil
import statistics
import textwrap
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "nir2-report.md"
ASSET_DIR = ROOT / "docs" / "nir2-assets"
DATASET_POINTER = ROOT / "datasets" / "latest.json"
DEFAULT_OUTPUT = Path.home() / "Downloads" / "Гарифуллин_ИИ_НИР-2_TeleBid.docx"
TOC_PAGE_MAP = ROOT / "docs" / "nir2-toc-pages.json"

BODY_FONT_SIZE = 11.5
BODY_LINE_SPACING = 1.15

FONT_REGULAR = "/System/Library/Fonts/Supplemental/Verdana.ttf"
FONT_BOLD = "/System/Library/Fonts/Supplemental/Verdana Bold.ttf"

INK = "172033"
MUTED = "52606D"
ACCENT = "E5A900"
LIGHT = "F3F5F7"
GRID = "C8CFD6"
POSITIVE = "2D7D46"
WARNING = "B46A00"


def pil_font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    return ImageFont.truetype(FONT_BOLD if bold else FONT_REGULAR, size=size)


def wrapped(draw: ImageDraw.ImageDraw, text: str, font: ImageFont.FreeTypeFont, width: int) -> list[str]:
    words = text.split()
    lines: list[str] = []
    current = ""
    for word in words:
        candidate = f"{current} {word}".strip()
        if draw.textbbox((0, 0), candidate, font=font)[2] <= width:
            current = candidate
        else:
            if current:
                lines.append(current)
            current = word
    if current:
        lines.append(current)
    return lines


def draw_centered_text(
    draw: ImageDraw.ImageDraw,
    box: tuple[int, int, int, int],
    text: str,
    font: ImageFont.FreeTypeFont,
    fill: str = "#172033",
    spacing: int = 8,
) -> None:
    left, top, right, bottom = box
    lines = wrapped(draw, text, font, right - left - 32)
    line_height = font.size + spacing
    y = top + ((bottom - top) - line_height * len(lines)) / 2
    for line in lines:
        bounds = draw.textbbox((0, 0), line, font=font)
        x = left + ((right - left) - (bounds[2] - bounds[0])) / 2
        draw.text((x, y), line, font=font, fill=fill)
        y += line_height


def draw_box(
    draw: ImageDraw.ImageDraw,
    box: tuple[int, int, int, int],
    title: str,
    subtitle: str = "",
    fill: str = "#FFFFFF",
    outline: str = "#172033",
) -> None:
    draw.rounded_rectangle(box, radius=22, fill=fill, outline=outline, width=3)
    left, top, right, bottom = box
    if subtitle:
        draw_centered_text(draw, (left + 10, top + 16, right - 10, top + 72), title, pil_font(25, True))
        draw_centered_text(
            draw,
            (left + 10, top + 70, right - 10, bottom - 12),
            subtitle,
            pil_font(18),
            fill="#52606D",
            spacing=5,
        )
    else:
        draw_centered_text(draw, box, title, pil_font(24, True))


def draw_arrow(
    draw: ImageDraw.ImageDraw,
    start: tuple[int, int],
    end: tuple[int, int],
    color: str = "#172033",
    width: int = 5,
    label: str | None = None,
) -> None:
    draw.line((start, end), fill=color, width=width)
    angle = math.atan2(end[1] - start[1], end[0] - start[0])
    size = 16
    for delta in (2.6, -2.6):
        point = (
            end[0] + size * math.cos(angle + delta),
            end[1] + size * math.sin(angle + delta),
        )
        draw.line((end, point), fill=color, width=width)
    if label:
        font = pil_font(17, True)
        bounds = draw.textbbox((0, 0), label, font=font)
        x = (start[0] + end[0]) / 2 - (bounds[2] - bounds[0]) / 2
        y = (start[1] + end[1]) / 2 - 26
        draw.rounded_rectangle(
            (x - 8, y - 4, x + bounds[2] - bounds[0] + 8, y + bounds[3] - bounds[1] + 6),
            radius=7,
            fill="#FFFFFF",
        )
        draw.text((x, y), label, font=font, fill=color)


def save_canvas(image: Image.Image, name: str) -> Path:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    path = ASSET_DIR / name
    image.save(path, format="PNG", optimize=True)
    return path


def build_notification_architecture() -> None:
    image = Image.new("RGB", (1800, 1000), "white")
    draw = ImageDraw.Draw(image)
    draw.text((70, 45), "ДВУХКАНАЛЬНАЯ ДОСТАВКА УВЕДОМЛЕНИЙ", font=pil_font(34, True), fill="#172033")
    draw.text(
        (70, 98),
        "Одна транзакционная причина, два независимых пути и восстановление по cursor",
        font=pil_font(21),
        fill="#52606D",
    )

    draw_box(draw, (70, 220, 390, 430), "Торговая команда", "ставка или закрытие аукциона", fill="#FFF8DE")
    draw_box(draw, (530, 220, 870, 430), "PostgreSQL", "событие + durable inbox в одной транзакции", fill="#EEF4FF")
    draw_box(draw, (1030, 160, 1400, 350), "Активный Mini App", "WebSocket live + causal gate", fill="#EAF7EF")
    draw_box(draw, (1030, 500, 1400, 690), "Telegram dispatcher", "status, attempts, retry", fill="#FFF3E8")
    draw_box(draw, (1480, 500, 1730, 690), "Чат бота", "deep link на торги", fill="#F2EDFF")
    draw_box(draw, (1030, 770, 1400, 930), "Повторно открытый Mini App", "GET/WS replay после sequence", fill="#EAF7EF")

    draw_arrow(draw, (390, 325), (530, 325), label="BEGIN / COMMIT")
    draw_arrow(draw, (870, 285), (1030, 255), color="#2D7D46", label="live")
    draw_arrow(draw, (870, 365), (1030, 575), color="#B46A00", label="outbox")
    draw_arrow(draw, (1400, 595), (1480, 595), color="#6B4BB8", label="Bot API")
    draw_arrow(draw, (870, 400), (1030, 845), color="#2D7D46", label="cursor replay")

    draw.rounded_rectangle((70, 530, 870, 890), radius=24, fill="#F7F8FA", outline="#C8CFD6", width=2)
    draw.text((110, 570), "Гарантии прикладного уровня", font=pil_font(25, True), fill="#172033")
    points = [
        "notificationId удаляет повторный показ",
        "sequence задаёт cursor inbox",
        "eventId связывает сообщение с причиной",
        "aggregateVersion запрещает показ поверх старого состояния",
        "FAILED остаётся в базе и подбирается worker после перезапуска",
    ]
    y = 630
    for point in points:
        draw.ellipse((112, y + 7, 124, y + 19), fill="#E5A900")
        for line in wrapped(draw, point, pil_font(19), 680):
            draw.text((145, y), line, font=pil_font(19), fill="#334155")
            y += 28
        y += 14
    save_canvas(image, "notification-architecture.png")


def build_stand_architecture() -> None:
    image = Image.new("RGB", (1800, 1050), "white")
    draw = ImageDraw.Draw(image)
    draw.text((70, 45), "ЭКСПЕРИМЕНТАЛЬНЫЙ СТЕНД TELEBID", font=pil_font(34, True), fill="#172033")
    draw.text((70, 98), "Команды отделены от наблюдаемого сетевого канала", font=pil_font(21), fill="#52606D")

    draw.rounded_rectangle((45, 155, 1755, 705), radius=28, fill="#FBFCFD", outline="#C8CFD6", width=3)
    draw.text((80, 180), "Docker project telebid-research", font=pil_font(21, True), fill="#52606D")

    draw_box(draw, (90, 290, 410, 500), "Генератор команд", "TypeScript, fixed seed, HTTP POST", fill="#FFF8DE")
    draw_box(draw, (560, 250, 930, 540), "NestJS API", "AuctionService\nSyncController\nWebSocket Gateway\nTelegram Dispatcher", fill="#EEF4FF")
    draw_box(draw, (1080, 250, 1420, 540), "PostgreSQL", "auctions · bids\nprocessed_commands\nauction_events\nnotifications", fill="#EAF7EF")
    draw_box(draw, (1500, 300, 1690, 490), "Telegram adapter", "fail first\nthen retry", fill="#FFF3E8")
    draw_box(draw, (90, 770, 430, 970), "Виртуальные клиенты", "polling · SSE · WebSocket\nMini App notification observer", fill="#F2EDFF")
    draw_box(draw, (620, 770, 950, 970), "Toxiproxy", "latency · jitter · disconnect", fill="#FFF3E8")
    draw_box(draw, (1150, 770, 1690, 970), "Датасет", "events · commands · notifications\nclients · trials · aggregates · manifest", fill="#EAF7EF")

    draw_arrow(draw, (410, 395), (560, 395), label="control API")
    draw_arrow(draw, (930, 395), (1080, 395), label="transaction")
    draw_arrow(draw, (1420, 395), (1500, 395), color="#B46A00", label="sendMessage")
    draw_arrow(draw, (430, 870), (620, 870), color="#6B4BB8", label="observer traffic")
    draw_arrow(draw, (950, 870), (780, 540), color="#6B4BB8", label="HTTP / SSE / WS")
    draw_arrow(draw, (930, 600), (1370, 770), color="#2D7D46", label="raw observations")
    draw_arrow(draw, (1080, 500), (1400, 770), color="#2D7D46", label="oracle")
    save_canvas(image, "stand-architecture.png")


def build_reconnect_sequence() -> None:
    image = Image.new("RGB", (1800, 1080), "white")
    draw = ImageDraw.Draw(image)
    draw.text((70, 40), "RECOVERY ПРИ ЗАКРЫТОМ MINI APP", font=pil_font(34, True), fill="#172033")
    participants = [
        (180, "Mini App"),
        (520, "Toxiproxy"),
        (870, "NestJS API"),
        (1220, "PostgreSQL"),
        (1580, "Telegram Bot"),
    ]
    for x, name in participants:
        draw.rounded_rectangle((x - 120, 115, x + 120, 185), radius=18, fill="#EEF4FF", outline="#172033", width=2)
        draw_centered_text(draw, (x - 115, 120, x + 115, 180), name, pil_font(19, True))
        draw.line((x, 185, x, 1010), fill="#AAB4BE", width=2)

    events = [
        (245, 180, 870, "subscribe(lastAppliedVersion)"),
        (315, 870, 1220, "eventsAfter(cursor)"),
        (385, 1220, 870, "ordered events"),
        (455, 870, 180, "sync complete + live"),
        (595, 870, 1220, "ставки + AUCTION_CLOSED + inbox"),
        (665, 1220, 1580, "pending notification"),
        (735, 1580, 1220, "first attempt FAILED"),
        (805, 1220, 1580, "retry → DELIVERED"),
        (905, 180, 870, "reconnect(version, sequence)"),
        (975, 870, 180, "state replay → notification replay"),
    ]
    draw.rounded_rectangle((45, 500, 690, 855), radius=22, fill="#FFF4F2", outline="#C84A3A", width=3)
    draw.text((80, 525), "FAULT", font=pil_font(24, True), fill="#9B2C24")
    draw.text((80, 565), "Mini App отключён", font=pil_font(21), fill="#9B2C24")
    draw.text((80, 605), "Торги продолжаются на сервере", font=pil_font(18), fill="#52606D")
    draw.text((80, 645), "Финальное событие live не приходит", font=pil_font(18), fill="#52606D")
    draw.text((80, 685), "Bot API временно ошибается", font=pil_font(18), fill="#52606D")
    draw.text((80, 725), "Записи остаются в PostgreSQL", font=pil_font(18), fill="#52606D")

    for y, start_x, end_x, label in events:
        if 500 <= y <= 855 and start_x == 180:
            continue
        color = "#2D7D46" if y >= 805 else ("#B46A00" if y >= 665 else "#172033")
        draw_arrow(draw, (start_x, y), (end_x, y), color=color, width=4, label=label)
    save_canvas(image, "reconnect-sequence.png")


def build_baseline_chart(summary: dict) -> None:
    trials = [trial for trial in summary["trials"] if trial["config"]["name"] == "baseline-live"]
    values: dict[tuple[str, str], float] = {}
    for kind in ("DIRECT", "REVERSE"):
        for transport in ("polling", "sse", "websocket"):
            p95s = [
                next(item for item in trial["transports"] if item["transport"] == transport)["p95LatencyMs"]
                for trial in trials
                if trial["config"]["auctionKind"] == kind
            ]
            values[(kind, transport)] = statistics.median(p95s)

    image = Image.new("RGB", (1800, 1000), "white")
    draw = ImageDraw.Draw(image)
    draw.text((70, 45), "LIVE-ДОСТАВКА В СТАБИЛЬНОМ ПРОФИЛЕ", font=pil_font(34, True), fill="#172033")
    draw.text((70, 98), "Медиана trial-p95, миллисекунды · меньше — лучше", font=pil_font(21), fill="#52606D")
    chart_left, chart_top, chart_right, chart_bottom = 150, 190, 1700, 820
    max_value = 300
    for tick in range(0, 301, 50):
        y = chart_bottom - (tick / max_value) * (chart_bottom - chart_top)
        draw.line((chart_left, y, chart_right, y), fill="#E1E5E9", width=2)
        draw.text((80, y - 13), str(tick), font=pil_font(17), fill="#52606D")

    colors = {"polling": "#E5A900", "sse": "#6B8AFD", "websocket": "#2D7D46"}
    labels = {"polling": "Polling", "sse": "SSE", "websocket": "WebSocket"}
    groups = [("DIRECT", "Прямой аукцион"), ("REVERSE", "Обратный аукцион")]
    group_centers = [570, 1270]
    bar_width = 120
    gap = 28
    for (kind, group_label), center in zip(groups, group_centers):
        positions = [center - bar_width - gap, center, center + bar_width + gap]
        for transport, x in zip(("polling", "sse", "websocket"), positions):
            value = values[(kind, transport)]
            height = (value / max_value) * (chart_bottom - chart_top)
            draw.rounded_rectangle(
                (x - bar_width / 2, chart_bottom - height, x + bar_width / 2, chart_bottom),
                radius=12,
                fill=colors[transport],
            )
            label = f"{int(value)}"
            bounds = draw.textbbox((0, 0), label, font=pil_font(22, True))
            draw.text((x - (bounds[2] - bounds[0]) / 2, chart_bottom - height - 38), label, font=pil_font(22, True), fill="#172033")
        bounds = draw.textbbox((0, 0), group_label, font=pil_font(22, True))
        draw.text((center - (bounds[2] - bounds[0]) / 2, 860), group_label, font=pil_font(22, True), fill="#172033")

    legend_x = 420
    for transport in ("polling", "sse", "websocket"):
        draw.rounded_rectangle((legend_x, 930, legend_x + 28, 958), radius=6, fill=colors[transport])
        draw.text((legend_x + 42, 928), labels[transport], font=pil_font(18), fill="#334155")
        legend_x += 300
    save_canvas(image, "baseline-latency.png")


def generate_assets() -> None:
    pointer = json.loads(DATASET_POINTER.read_text(encoding="utf-8"))
    summary_path = Path(pointer["outputDirectory"]) / "summary.json"
    summary = json.loads(summary_path.read_text(encoding="utf-8"))
    build_notification_architecture()
    build_stand_architecture()
    build_reconnect_sequence()
    build_baseline_chart(summary)


def set_run_font(
    run,
    name: str = "Times New Roman",
    size: float | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
    color: str | None = None,
) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def shade_cell(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for margin_name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color: str = GRID, size: int = 6) -> None:
    properties = table._tbl.tblPr
    borders = properties.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        properties.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), str(size))
        element.set(qn("w:color"), color)


def table_widths(rows: list[list[str]]) -> list[int]:
    column_count = len(rows[0])
    maxima = []
    for index in range(column_count):
        lengths = [min(80, max(4, len(row[index]))) for row in rows]
        maxima.append(max(lengths[0] * 1.25, statistics.mean(lengths), max(lengths) * 0.55))
    if column_count == 2:
        maxima[0] = min(maxima[0], sum(maxima) * 0.42)
    minimum = 0.11 if column_count >= 5 else 0.16
    total = sum(maxima)
    shares = [max(minimum, value / total) for value in maxima]
    share_total = sum(shares)
    widths = [round(9360 * share / share_total) for share in shares]
    widths[-1] += 9360 - sum(widths)
    return widths


def set_table_geometry(table, widths: list[int]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    properties = table._tbl.tblPr
    table_width = properties.find(qn("w:tblW"))
    if table_width is None:
        table_width = OxmlElement("w:tblW")
        properties.append(table_width)
    table_width.set(qn("w:w"), "9360")
    table_width.set(qn("w:type"), "dxa")
    indent = properties.find(qn("w:tblInd"))
    if indent is None:
        indent = OxmlElement("w:tblInd")
        properties.append(indent)
    indent.set(qn("w:w"), "120")
    indent.set(qn("w:type"), "dxa")
    layout = properties.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        properties.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(width))
        grid.append(column)

    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = Inches(width / 1440)
            properties = cell._tc.get_or_add_tcPr()
            tc_width = properties.find(qn("w:tcW"))
            if tc_width is None:
                tc_width = OxmlElement("w:tcW")
                properties.append(tc_width)
            tc_width.set(qn("w:w"), str(width))
            tc_width.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    properties = row._tr.get_or_add_trPr()
    marker = OxmlElement("w:tblHeader")
    marker.set(qn("w:val"), "true")
    properties.append(marker)


def add_page_field(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "2"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run = paragraph.add_run()
    run._r.extend([begin, instruction, separate, text, end])
    set_run_font(run, size=11)


def extract_toc_entries(source: str) -> list[tuple[int, str]]:
    entries: list[tuple[int, str]] = []
    for line in source.splitlines():
        heading = re.match(r"^(#{1,2})\s+(.+)$", line.strip())
        if heading:
            entries.append((len(heading.group(1)), heading.group(2)))
    return entries


def add_toc(doc: Document, source: str) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(8)
    title_run = title.add_run("СОДЕРЖАНИЕ")
    set_run_font(title_run, size=14, bold=True)

    page_map = json.loads(TOC_PAGE_MAP.read_text(encoding="utf-8")) if TOC_PAGE_MAP.exists() else {}
    for level, text in extract_toc_entries(source):
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.left_indent = Cm(0.65 if level == 2 else 0)
        paragraph.paragraph_format.right_indent = Cm(0)
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0
        paragraph.paragraph_format.tab_stops.add_tab_stop(
            Cm(15.8),
            WD_TAB_ALIGNMENT.RIGHT,
            WD_TAB_LEADER.DOTS,
        )
        name = paragraph.add_run(text)
        set_run_font(name, size=9.6, bold=level == 1)
        page = paragraph.add_run(f"\t{page_map.get(text, '00')}")
        set_run_font(page, size=9.6, bold=level == 1)


def add_title_page(doc: Document) -> None:
    def centered(text: str, size: float = 12.5, bold: bool = False, after: float = 0) -> None:
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(after)
        paragraph.paragraph_format.line_spacing = 1.15
        run = paragraph.add_run(text)
        set_run_font(run, size=size, bold=bold)

    centered("Министерство науки и высшего образования Российской Федерации", 12, False, 2)
    centered("ФЕДЕРАЛЬНОЕ ГОСУДАРСТВЕННОЕ АВТОНОМНОЕ", 12, True)
    centered("ОБРАЗОВАТЕЛЬНОЕ УЧРЕЖДЕНИЕ ВЫСШЕГО ОБРАЗОВАНИЯ", 12, True)
    centered("«НАЦИОНАЛЬНЫЙ ИССЛЕДОВАТЕЛЬСКИЙ УНИВЕРСИТЕТ ИТМО»", 12, True, 4)
    centered("(Университет ИТМО)", 12, False, 18)
    centered("Факультет программной инженерии и компьютерной техники", 12, False, 3)
    centered("Образовательная программа «Веб-технологии»", 12, False, 3)
    centered("Направление подготовки 09.04.04 «Программная инженерия»", 12, False, 34)

    centered("ОТЧЁТ", 15, True, 4)
    centered("по научно-исследовательской работе", 14, True, 22)
    centered(
        "Тема задания: «Исследование архитектурных решений обработки конкурентных торгов и синхронизации состояния в веб-платформах прямых и обратных аукционов»",
        13,
        True,
        28,
    )

    details = [
        "Обучающийся: Гарифуллин Искандар Ильданович, 506911",
        "Руководитель практики от университета: Государев Илья Борисович",
        "Практика пройдена с оценкой ____________________",
        "Дата: 30.06.2026",
    ]
    for text in details:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.left_indent = Cm(7.0)
        paragraph.paragraph_format.space_after = Pt(6)
        run = paragraph.add_run(text)
        set_run_font(run, size=12.5)

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(48)
    centered("Санкт-Петербург", 12.5)
    centered("2026", 12.5)


def add_inline_runs(paragraph, text: str, size: float = BODY_FONT_SIZE) -> None:
    parts = re.split(r"(`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_run_font(run, name="Courier New", size=max(9.5, size - 1.5))
        else:
            run = paragraph.add_run(part)
            set_run_font(run, size=size)


def add_body_paragraph(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="Normal")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.first_line_indent = Cm(1.25)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = BODY_LINE_SPACING
    add_inline_runs(paragraph, text)


def add_caption(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(5)
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run(text)
    set_run_font(run, size=9.5)


def add_figure(doc: Document, filename: str, caption: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run()
    shape = run.add_picture(str(ASSET_DIR / filename), width=Inches(5.9))
    shape._inline.docPr.set("title", caption)
    shape._inline.docPr.set("descr", caption)
    add_caption(doc, caption)


def create_numbering(doc: Document, kind: str) -> int:
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(node.get(qn("w:abstractNumId"))) for node in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    number_format = OxmlElement("w:numFmt")
    number_format.set(qn("w:val"), "decimal" if kind == "decimal" else "bullet")
    level.append(number_format)
    level_text = OxmlElement("w:lvlText")
    level_text.set(qn("w:val"), "%1." if kind == "decimal" else "•")
    level.append(level_text)
    justification = OxmlElement("w:lvlJc")
    justification.set(qn("w:val"), "left")
    level.append(justification)
    paragraph_properties = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "720")
    tabs.append(tab)
    paragraph_properties.append(tabs)
    indentation = OxmlElement("w:ind")
    indentation.set(qn("w:left"), "720")
    indentation.set(qn("w:hanging"), "360")
    paragraph_properties.append(indentation)
    level.append(paragraph_properties)
    abstract.append(level)
    numbering.append(abstract)

    number = OxmlElement("w:num")
    number.set(qn("w:numId"), str(num_id))
    reference = OxmlElement("w:abstractNumId")
    reference.set(qn("w:val"), str(abstract_id))
    number.append(reference)
    numbering.append(number)
    return num_id


def add_list_paragraph(
    doc: Document,
    text: str,
    num_id: int,
    *,
    size: float = BODY_FONT_SIZE,
    line_spacing: float = BODY_LINE_SPACING,
) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = line_spacing
    properties = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    level = OxmlElement("w:ilvl")
    level.set(qn("w:val"), "0")
    number = OxmlElement("w:numId")
    number.set(qn("w:val"), str(num_id))
    num_pr.extend([level, number])
    properties.append(num_pr)
    add_inline_runs(paragraph, text, size=size)


def add_markdown_table(doc: Document, rows: list[list[str]]) -> None:
    widths = table_widths(rows)
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    set_table_geometry(table, widths)
    set_table_borders(table)
    set_repeat_table_header(table.rows[0])
    for row_index, row in enumerate(rows):
        for column_index, value in enumerate(row):
            cell = table.cell(row_index, column_index)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            if row_index == 0:
                shade_cell(cell, "E8EEF5")
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.05
            short_value = len(value) < 18 and not re.search(r"[а-яА-Я]{12,}", value)
            paragraph.alignment = (
                WD_ALIGN_PARAGRAPH.CENTER
                if row_index == 0 or short_value or re.fullmatch(r"[\d.,–— %]+", value)
                else WD_ALIGN_PARAGRAPH.LEFT
            )
            add_inline_runs(paragraph, value, size=8.8)
            for run in paragraph.runs:
                if row_index == 0:
                    run.bold = True
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(2)


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    normal.font.size = Pt(BODY_FONT_SIZE)
    normal.paragraph_format.line_spacing = BODY_LINE_SPACING
    normal.paragraph_format.space_after = Pt(0)

    heading1 = styles["Heading 1"]
    heading1.font.name = "Times New Roman"
    heading1._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    heading1._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    heading1.font.size = Pt(13.5)
    heading1.font.bold = True
    heading1.font.color.rgb = RGBColor(0, 0, 0)
    heading1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading1.paragraph_format.space_before = Pt(0)
    heading1.paragraph_format.space_after = Pt(8)
    heading1.paragraph_format.keep_with_next = True

    heading2 = styles["Heading 2"]
    heading2.font.name = "Times New Roman"
    heading2._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    heading2._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    heading2.font.size = Pt(12)
    heading2.font.bold = True
    heading2.font.color.rgb = RGBColor(0, 0, 0)
    heading2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    heading2.paragraph_format.space_before = Pt(6)
    heading2.paragraph_format.space_after = Pt(2)
    heading2.paragraph_format.keep_with_next = True


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    raw: list[str] = []
    index = start
    while index < len(lines) and lines[index].strip().startswith("|"):
        raw.append(lines[index].strip())
        index += 1
    rows = [[cell.strip() for cell in row.strip("|").split("|")] for row in raw]
    if len(rows) >= 2 and all(re.fullmatch(r":?-{3,}:?", cell) for cell in rows[1]):
        rows.pop(1)
    return rows, index


def add_report_body(doc: Document, source: str) -> None:
    lines = source.splitlines()
    index = 0
    first_heading = True
    current_heading1 = ""
    active_list: str | None = None
    list_num_id: int | None = None
    while index < len(lines):
        line = lines[index].strip()
        if not line:
            active_list = None
            list_num_id = None
            index += 1
            continue

        figure = re.fullmatch(r"\[\[FIGURE:([^|]+)\|(.+)\]\]", line)
        if figure:
            add_figure(doc, figure.group(1), figure.group(2))
            active_list = None
            index += 1
            continue

        if line.startswith("|"):
            rows, index = parse_table(lines, index)
            add_markdown_table(doc, rows)
            active_list = None
            continue

        heading = re.match(r"^(#{1,2})\s+(.+)$", line)
        if heading:
            level = len(heading.group(1))
            text = heading.group(2)
            page_break_before = level == 1 and not first_heading and not text.startswith("ПРИЛОЖЕНИЕ Б")
            paragraph = doc.add_paragraph(style=f"Heading {level}")
            paragraph.paragraph_format.keep_with_next = True
            paragraph.paragraph_format.page_break_before = page_break_before
            run = paragraph.add_run(text)
            set_run_font(run, size=13.5 if level == 1 else 12, bold=True)
            if level == 1:
                current_heading1 = text
            first_heading = False
            active_list = None
            list_num_id = None
            index += 1
            continue

        list_match = re.match(r"^(\d+)\.\s+(.+)$", line)
        bullet_match = re.match(r"^-\s+(.+)$", line)
        if list_match or bullet_match:
            kind = "decimal" if list_match else "bullet"
            text = list_match.group(2) if list_match else bullet_match.group(1)
            if active_list != kind or list_num_id is None:
                list_num_id = create_numbering(doc, kind)
                active_list = kind
            is_references = current_heading1 == "СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ"
            add_list_paragraph(
                doc,
                text,
                list_num_id,
                size=10 if is_references else BODY_FONT_SIZE,
                line_spacing=1.0 if is_references else BODY_LINE_SPACING,
            )
            index += 1
            continue

        if re.match(r"^(Таблица|Рисунок)\s+\d+", line):
            add_caption(doc, line)
            index += 1
            continue

        paragraph_lines = [line]
        index += 1
        while index < len(lines):
            candidate = lines[index].strip()
            if not candidate:
                break
            if (
                candidate.startswith("#")
                or candidate.startswith("|")
                or candidate.startswith("[[FIGURE:")
                or re.match(r"^\d+\.\s+", candidate)
                or candidate.startswith("- ")
            ):
                break
            paragraph_lines.append(candidate)
            index += 1
        add_body_paragraph(doc, " ".join(paragraph_lines))
        active_list = None
        list_num_id = None


def set_update_fields(doc: Document) -> None:
    settings = doc.settings._element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")


def build_document(output: Path, page_count: int) -> None:
    generate_assets()
    source = SOURCE.read_text(encoding="utf-8").replace("{{PAGE_COUNT}}", str(page_count))
    doc = Document()
    configure_styles(doc)
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(1.5)
    section.header_distance = Cm(1.0)
    section.footer_distance = Cm(1.0)
    section.different_first_page_header_footer = True
    add_page_field(section.footer.paragraphs[0])

    doc.core_properties.title = "НИР-2 TeleBid"
    doc.core_properties.subject = "Архитектурные решения конкурентных торгов и синхронизации состояния"
    doc.core_properties.author = "Гарифуллин Искандар Ильданович"
    doc.core_properties.keywords = "Telegram Mini App, WebSocket, SSE, polling, аукцион, НИР"

    add_title_page(doc)
    doc.add_page_break()
    add_toc(doc, source)
    doc.add_page_break()
    add_report_body(doc, source)
    set_update_fields(doc)

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--output", type=Path, default=DEFAULT_OUTPUT)
    parser.add_argument("--page-count", type=int, default=25)
    parser.add_argument("--copy-to", type=Path)
    args = parser.parse_args()
    build_document(args.output, args.page_count)
    if args.copy_to:
        args.copy_to.parent.mkdir(parents=True, exist_ok=True)
        shutil.copy2(args.output, args.copy_to)
    print(args.output)


if __name__ == "__main__":
    main()
